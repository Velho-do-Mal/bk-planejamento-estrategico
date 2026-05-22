"""
views.py — BK Planejamento Estratégico
CORREÇÕES APLICADAS:
  1. save_meta: lookup por NOME (não por índice) → preserva dados ao editar/reordenar KPIs
  2. _normalize_partner_rows / _normalize_area_rows: respeita flag 'excluir'
  3. fig_swot_quadrant: texto branco sobre fundo colorido (visibilidade)
  4. kpis view (ex-s): nome descritivo, chave duplicada removida
  5. logout: protegido por POST (já em views.py de accounts)
  6. build_word_report: geração de relatório Word com gráficos matplotlib
  7. export_word_view: endpoint de download do Word
"""

import copy
import io
import json
import zipfile
from datetime import date, datetime
from typing import List, Optional

import numpy as np
import pandas as pd
import plotly.graph_objs as go
import plotly.express as px

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render

from .models import PlanningData


# ─── Paleta BK ──────────────────────────────────────────────────────────────
BK_BLUE       = "#1565C0"
BK_BLUE_LIGHT = "#42A5F5"
BK_TEAL       = "#00897B"
BK_GREEN      = "#43A047"
BK_ORANGE     = "#FB8C00"
BK_RED        = "#E53935"
BK_PURPLE     = "#7B1FA2"
BK_GRAY       = "#546E7A"

SWOT_COLORS = {
    "Força":        "#43A047",
    "Fraqueza":     "#E53935",
    "Oportunidade": "#1565C0",
    "Ameaça":       "#FB8C00",
}
STATUS_COLORS = {
    "Concluído":    BK_GREEN,
    "Em andamento": BK_ORANGE,
    "Pendente":     BK_GRAY,
    "Atrasado":     BK_RED,
}


# ─── Helpers de dados ────────────────────────────────────────────────────────

def get_planning() -> dict:
    obj = PlanningData.get_or_create_default()
    dados = obj.dados or {}
    dados.setdefault("partners", [])
    dados.setdefault("areas",    [])
    dados.setdefault("swot",     [])
    dados.setdefault("actions",  [])
    dados.setdefault("strategic", {
        "visao": "", "missao": "", "valores": "", "posicionamento": "",
        "proposta_valor": "", "publico_alvo": "", "diferenciais": "",
        "pilares": "", "objetivos_estrategicos": "", "notas": "",
    })
    # Migrar chave legada 'okrs' → 's'
    if "okrs" in dados and dados["okrs"]:
        if "s" not in dados or not dados["s"]:
            dados["s"] = dados["okrs"]
        del dados["okrs"]
        obj.dados = dados
        obj.save()
    dados.setdefault("s", [])
    return dados


def save_planning(dados: dict):
    dados_copia = copy.deepcopy(dados)
    nz = sum(
        1 for o in dados_copia.get("s", [])
        for m in o.get("meses", [])
        if (m.get("previsto") or 0) != 0 or (m.get("realizado") or 0) != 0
    )
    print(f"[save_planning] kpis={len(dados_copia.get('s', []))} | non_zero_meses={nz}", flush=True)
    updated = PlanningData.objects.filter(slug="bk").update(dados=dados_copia)
    if not updated:
        PlanningData.objects.create(slug="bk", dados=dados_copia)
    print(f"[save_planning] rows_updated={updated}", flush=True)


def _clean_text(value, default=""):
    return str(value if value is not None else default).strip()


def _normalize_partner_rows(rows: list) -> list:
    """Normaliza linhas de sócios, respeitando flag 'excluir'."""
    normalized = []
    for row in rows:
        if row.get("excluir"):          # BUG FIX: respeitar checkbox de exclusão
            continue
        nome = _clean_text(row.get("nome"))
        if not nome:
            continue
        normalized.append({
            "nome":        nome,
            "cargo":       _clean_text(row.get("cargo")),
            "email":       _clean_text(row.get("email")),
            "telefone":    _clean_text(row.get("telefone")),
            "observacoes": _clean_text(row.get("observacoes")),
        })
    return normalized


def _normalize_area_rows(rows: list) -> list:
    """Normaliza linhas de áreas, respeitando flag 'excluir'."""
    normalized = []
    for row in rows:
        if row.get("excluir"):          # BUG FIX: respeitar checkbox de exclusão
            continue
        area = _clean_text(row.get("area"))
        if not area:
            continue
        normalized.append({
            "area":        area,
            "responsavel": _clean_text(row.get("responsavel")),
            "email":       _clean_text(row.get("email")),
            "observacoes": _clean_text(row.get("observacoes")),
        })
    return normalized


def _group_swot_items(swot_items: list) -> dict:
    grupos = {"Força": [], "Fraqueza": [], "Oportunidade": [], "Ameaça": []}
    for item in swot_items or []:
        tipo      = _clean_text(item.get("tipo"))
        descricao = _clean_text(item.get("descricao"))
        prioridade = _clean_text(item.get("prioridade"), "Média")
        if tipo in grupos and descricao:
            grupos[tipo].append({"tipo": tipo, "descricao": descricao,
                                 "prioridade": prioridade or "Média"})
    return grupos


def _safe_date(s) -> Optional[date]:
    if not s:
        return None
    try:
        if isinstance(s, date):
            return s
        return datetime.strptime(str(s)[:10], "%Y-%m-%d").date()
    except Exception:
        return None


def _ensure_okr_meses(okr: dict) -> dict:
    """Garante lista de 36 meses com previsto/realizado como float."""
    try:
        if not isinstance(okr, dict):
            okr = {}
        meses = okr.get("meses", [])
        if not isinstance(meses, list):
            meses = []
        meses_corrigidos = []
        for i in range(36):
            mes = meses[i] if i < len(meses) and isinstance(meses[i], dict) else {}
            try:
                previsto  = float(mes.get("previsto",  0) or 0)
            except (TypeError, ValueError):
                previsto  = 0.0
            try:
                realizado = float(mes.get("realizado", 0) or 0)
            except (TypeError, ValueError):
                realizado = 0.0
            meses_corrigidos.append({"previsto": previsto, "realizado": realizado})
        okr["meses"] = meses_corrigidos
        return okr
    except Exception:
        return {"meses": [{"previsto": 0.0, "realizado": 0.0} for _ in range(36)]}


def _month_labels_for_okr(okr: dict) -> List[str]:
    inicio_str = okr.get("inicio", "")
    try:
        dt = datetime.strptime(str(inicio_str)[:7], "%Y-%m") if inicio_str else datetime(date.today().year, 1, 1)
    except Exception:
        dt = datetime(date.today().year, 1, 1)
    labels = []
    for i in range(36):
        m = (dt.month - 1 + i) % 12 + 1
        y = dt.year + (dt.month - 1 + i) // 12
        labels.append(f"{m:02d}/{y}")
    return labels


# ─── Gráficos Plotly ─────────────────────────────────────────────────────────

def _fig_layout(fig, title="", height=380):
    fig.update_layout(
        title=title, height=height,
        margin=dict(l=40, r=20, t=40 if title else 20, b=40),
        paper_bgcolor="white", plot_bgcolor="#F8FAFC",
        font=dict(family="Segoe UI, sans-serif", size=12),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
    )
    return fig


def fig_okr_monthly(okr: dict) -> str:
    okr    = _ensure_okr_meses(okr)
    labels = _month_labels_for_okr(okr)
    prev   = [float(m.get("previsto",  0)) for m in okr["meses"]]
    real   = [float(m.get("realizado", 0)) for m in okr["meses"]]
    unidade = okr.get("unidade", "")
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Planejado", x=labels, y=prev,
                         marker_color=BK_BLUE_LIGHT, opacity=0.7))
    fig.add_trace(go.Scatter(name="Realizado", x=labels, y=real,
                             mode="lines+markers",
                             line=dict(color=BK_GREEN, width=2.5),
                             marker=dict(size=6)))
    _fig_layout(fig, f"Mensal — {okr.get('nome','')} ({unidade})", height=360)
    return fig.to_json()


def fig_okr_cumulative(okr: dict) -> str:
    okr      = _ensure_okr_meses(okr)
    labels   = _month_labels_for_okr(okr)
    prev     = [float(m.get("previsto",  0)) for m in okr["meses"]]
    real     = [float(m.get("realizado", 0)) for m in okr["meses"]]
    cum_prev = list(np.cumsum(prev))
    cum_real = list(np.cumsum(real))
    fig = go.Figure()
    fig.add_trace(go.Scatter(name="Acumulado Planejado", x=labels, y=cum_prev,
                             mode="lines", line=dict(color=BK_BLUE, dash="dash", width=2)))
    fig.add_trace(go.Scatter(name="Acumulado Realizado", x=labels, y=cum_real,
                             mode="lines+markers", line=dict(color=BK_GREEN, width=2.5),
                             fill="tozeroy", fillcolor="rgba(67,160,71,0.08)"))
    _fig_layout(fig, f"Acumulado — {okr.get('nome','')}", height=320)
    return fig.to_json()


def fig_okr_gauge(okr: dict) -> str:
    okr   = _ensure_okr_meses(okr)
    tp    = sum(float(m.get("previsto",  0)) for m in okr["meses"])
    tr    = sum(float(m.get("realizado", 0)) for m in okr["meses"])
    pct   = (tr / tp * 100) if tp > 0 else 0
    color = BK_GREEN if pct >= 90 else (BK_ORANGE if pct >= 70 else BK_RED)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=pct,
        number={"suffix": "%", "font": {"size": 22}},
        title={"text": okr.get("nome", "")[:25], "font": {"size": 11}},
        gauge={
            "axis":  {"range": [0, 150], "tickwidth": 1},
            "bar":   {"color": color},
            "steps": [
                {"range": [0,   70], "color": "#FEE2E2"},
                {"range": [70,  90], "color": "#FEF3C7"},
                {"range": [90, 150], "color": "#D1FAE5"},
            ],
            "threshold": {"line": {"color": BK_BLUE, "width": 3}, "value": 100},
        }
    ))
    fig.update_layout(height=220, margin=dict(l=20, r=20, t=40, b=10),
                      paper_bgcolor="white", font=dict(family="Segoe UI"))
    return fig.to_json()


def fig_swot_quadrant(swot_items: list) -> str:
    """
    BUG FIX: labels dos quadrantes agora usam texto BRANCO sobre fundo colorido
    (fg) em vez de texto escuro sobre fundo claro (confundia com o quadrante).
    """
    fig = go.Figure()
    quadrants = {
        "Força":        (0.25, 0.75, "#D1FAE5", "#065F46"),
        "Oportunidade": (0.75, 0.75, "#DBEAFE", "#1E3A8A"),
        "Fraqueza":     (0.25, 0.25, "#FEE2E2", "#991B1B"),
        "Ameaça":       (0.75, 0.25, "#FEF3C7", "#92400E"),
    }

    for tipo, (cx, cy, bg, fg) in quadrants.items():
        x0 = 0 if cx < 0.5 else 0.5
        x1 = 0.5 if cx < 0.5 else 1
        y0 = 0 if cy < 0.5 else 0.5
        y1 = 0.5 if cy < 0.5 else 1
        fig.add_shape(type="rect", x0=x0, x1=x1, y0=y0, y1=y1,
                      xref="paper", yref="paper",
                      fillcolor=bg, line=dict(color="#CBD5E1", width=2))
        # TEXTO BRANCO sobre fundo colorido escuro → sempre visível
        fig.add_annotation(
            x=cx, y=cy + 0.16,
            xref="paper", yref="paper",
            text=f"<b>{tipo}</b>",
            showarrow=False,
            font=dict(size=14, color="white"),
            bgcolor=fg,
            bordercolor=fg,
            borderwidth=2,
            borderpad=5,
            opacity=0.95,
        )

    for tipo, (cx, cy, bg, fg) in quadrants.items():
        items = [s for s in swot_items if s.get("tipo") == tipo]
        for i, item in enumerate(items):
            jitter_x = (i % 3 - 1) * 0.07
            jitter_y = -(i // 3) * 0.09
            y_pos = cy - 0.06 + jitter_y
            fig.add_trace(go.Scatter(
                x=[cx + jitter_x], y=[y_pos],
                mode="markers+text",
                marker=dict(size=18, color=SWOT_COLORS.get(tipo, BK_GRAY),
                            line=dict(color="white", width=2)),
                text=[item.get("prioridade", "")[:1]],
                textfont=dict(color="white", size=10, family="Segoe UI Bold"),
                textposition="middle center",
                hovertext=f"<b>{item.get('prioridade','')}</b><br>{item.get('descricao','')}",
                hoverinfo="text",
                name=tipo,
                showlegend=(i == 0),
            ))

    fig.update_layout(
        height=430,
        paper_bgcolor="white",
        xaxis=dict(showticklabels=False, showgrid=False, zeroline=False, range=[0, 1]),
        yaxis=dict(showticklabels=False, showgrid=False, zeroline=False, range=[0, 1]),
        margin=dict(l=10, r=10, t=30, b=10),
        title=dict(text="Matriz SWOT", font=dict(size=14, color=BK_BLUE)),
        font=dict(family="Segoe UI"),
        hoverlabel=dict(bgcolor="white", font_size=12),
    )
    return fig.to_json()


def fig_actions_status(dados: dict) -> Optional[str]:
    actions = dados.get("actions", [])
    if not actions:
        return None
    counts = {}
    for a in actions:
        status_val = a.get("status", "Pendente")
        counts[status_val] = counts.get(status_val, 0) + 1
    fig = go.Figure(go.Pie(
        labels=list(counts.keys()), values=list(counts.values()), hole=0.5,
        marker=dict(colors=[STATUS_COLORS.get(k, BK_GRAY) for k in counts.keys()]),
        textinfo="label+percent",
    ))
    _fig_layout(fig, "Status dos Planos de Ação", height=320)
    return fig.to_json()


def fig_actions_timeline(dados: dict) -> Optional[str]:
    actions = dados.get("actions", [])
    today   = date.today()
    rows    = []
    for a in actions:
        d_ini = _safe_date(a.get("data_inicio"))    or today
        d_fim = _safe_date(a.get("data_vencimento")) or today
        if d_ini > d_fim:
            d_fim = d_ini
        rows.append({
            "Tarefa":      a.get("titulo", "")[:35],
            "Início":      d_ini,
            "Fim":         d_fim,
            "Status":      a.get("status", "Pendente"),
            "Responsável": a.get("responsavel", ""),
        })
    if not rows:
        return None
    df  = pd.DataFrame(rows).sort_values("Início")
    fig = px.timeline(df, x_start="Início", x_end="Fim", y="Tarefa",
                      color="Status", hover_data=["Responsável"],
                      color_discrete_map=STATUS_COLORS)
    fig.update_yaxes(autorange="reversed")
    _fig_layout(fig, "Linha do Tempo — Planos de Ação",
                height=max(300, len(rows) * 35 + 80))
    return fig.to_json()


def fig_okrs_overview(dados: dict) -> Optional[str]:
    okrs = dados.get("s", [])
    if not okrs:
        return None
    names, prevs, reals, pcts = [], [], [], []
    for o in okrs:
        o   = _ensure_okr_meses(o)
        tp  = sum(float(m.get("previsto",  0)) for m in o["meses"])
        tr  = sum(float(m.get("realizado", 0)) for m in o["meses"])
        pct = (tr / tp * 100) if tp > 0 else 0
        names.append(o.get("nome", "")[:30])
        prevs.append(tp); reals.append(tr); pcts.append(pct)
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Planejado", x=names, y=prevs,
                         marker_color=BK_BLUE_LIGHT, opacity=0.7))
    fig.add_trace(go.Bar(name="Realizado", x=names, y=reals,
                         marker_color=BK_GREEN))
    fig.add_trace(go.Scatter(name="% Realização", x=names, y=pcts,
                             mode="markers+text", yaxis="y2",
                             marker=dict(size=10, color=BK_ORANGE),
                             text=[f"{p:.0f}%" for p in pcts],
                             textposition="top center"))
    fig.update_layout(
        barmode="group", height=380,
        yaxis2=dict(overlaying="y", side="right",
                    title="% Realização", range=[0, 160]),
        paper_bgcolor="white", plot_bgcolor="#F8FAFC",
        margin=dict(l=40, r=60, t=40, b=60),
        font=dict(family="Segoe UI"),
        title="Visão Geral KPIs — Planejado vs Realizado",
    )
    return fig.to_json()


# ─── Exportações ─────────────────────────────────────────────────────────────

def export_excel(dados: dict) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if dados.get("partners"):
            pd.DataFrame(dados["partners"]).to_excel(writer, sheet_name="Socios_Gestores", index=False)
        if dados.get("areas"):
            pd.DataFrame(dados["areas"]).to_excel(writer, sheet_name="Areas", index=False)
        if dados.get("swot"):
            pd.DataFrame(dados["swot"]).to_excel(writer, sheet_name="SWOT", index=False)
        if dados.get("s"):
            rows = []
            for o in dados["s"]:
                o   = _ensure_okr_meses(o)
                row = {"nome": o.get("nome"), "area": o.get("area"),
                       "unidade": o.get("unidade"), "inicio": o.get("inicio")}
                for i, m in enumerate(o["meses"]):
                    row[f"M{i+1:02d}_prev"] = m.get("previsto",  0)
                    row[f"M{i+1:02d}_real"] = m.get("realizado", 0)
                rows.append(row)
            pd.DataFrame(rows).to_excel(writer, sheet_name="KPIs", index=False)
        if dados.get("actions"):
            pd.DataFrame(dados["actions"]).to_excel(writer, sheet_name="Planos_Acao", index=False)
        pd.DataFrame([dados.get("strategic", {})]).to_excel(writer, sheet_name="Estrategia", index=False)
    return output.getvalue()


def export_csv_zip(dados: dict) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for key, name in [("partners","socios.csv"),("areas","areas.csv"),
                           ("swot","swot.csv"),("actions","planos_acao.csv")]:
            if dados.get(key):
                zf.writestr(name, pd.DataFrame(dados[key]).to_csv(index=False))
    return buf.getvalue()


def _create_kpi_chart_image(okr: dict) -> bytes:
    """
    Gera imagem PNG de um KPI para o relatório Word.
    Planejado = LINHA (meta), Realizado = BARRAS VERTICAIS.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    okr    = _ensure_okr_meses(okr)
    labels = _month_labels_for_okr(okr)
    prev_v = [float(m.get("previsto",  0)) for m in okr["meses"]]
    real_v = [float(m.get("realizado", 0)) for m in okr["meses"]]

    # Mostrar meses com dados + mínimo de 12
    last_nz = 11
    for i in range(35, -1, -1):
        if prev_v[i] != 0 or real_v[i] != 0:
            last_nz = i
            break
    show_n  = max(last_nz + 1, 12)
    labels  = labels[:show_n]
    prev_v  = prev_v[:show_n]
    real_v  = real_v[:show_n]

    fig, ax = plt.subplots(figsize=(14, 3.8))
    x       = np.arange(len(labels))

    # Barras verticais → Realizado
    ax.bar(x, real_v, color="#43A047", alpha=0.85, label="Realizado",
           width=0.65, zorder=3)
    # Linha → Planejado (meta)
    ax.plot(x, prev_v, color="#1565C0", linewidth=2.5, marker="o",
            markersize=5, label="Planejado (Meta)", zorder=4)

    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel(okr.get("unidade", ""), fontsize=9)
    ax.legend(fontsize=9, loc="upper left")
    ax.grid(axis="y", alpha=0.25, zorder=0)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def build_word_report(dados: dict) -> bytes:
    """Gera relatório .docx completo com gráficos matplotlib."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc      = Document()
    today_str = date.today().strftime("%d/%m/%Y")
    strategic = dados.get("strategic", {})
    okrs      = dados.get("s", [])
    actions   = dados.get("actions", [])
    swot_list = dados.get("swot", [])

    C_BLUE  = RGBColor(21, 101, 192)
    C_GREEN = RGBColor(67, 160, 71)
    C_GRAY  = RGBColor(84, 110, 122)
    C_WHITE = RGBColor(255, 255, 255)
    C_RED   = RGBColor(229, 57, 53)
    C_ORG   = RGBColor(251, 140, 0)

    # ── Margens ──────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def _shd(cell, hex_fill):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"),  hex_fill)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"),   "clear")
        tcPr.append(shd)

    def _header_row(table, headers, fill="1565C0"):
        row = table.rows[0]
        for j, h in enumerate(headers):
            cell = row.cells[j]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = C_WHITE
            _shd(cell, fill)

    def _color_pct(pct):
        return C_GREEN if pct >= 95 else (C_ORG if pct >= 70 else C_RED)

    # ── Capa ─────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    h = doc.add_heading("Planejamento Estratégico", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = C_BLUE

    sub = doc.add_paragraph("BK Engenharia e Tecnologia")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.runs[0]; r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = C_BLUE

    d = doc.add_paragraph(f"Gerado em {today_str}  |  Horizonte: 36 meses")
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── Norte Estratégico ────────────────────────────────────
    h1 = doc.add_heading("Norte Estratégico", 1)
    for run in h1.runs: run.font.color.rgb = C_BLUE

    fields = [
        ("Visão",               strategic.get("visao")               or "—"),
        ("Missão",              strategic.get("missao")              or "—"),
        ("Valores",             strategic.get("valores")             or "—"),
        ("Posicionamento",      strategic.get("posicionamento")      or "—"),
        ("Proposta de Valor",   strategic.get("proposta_valor")      or "—"),
        ("Público-Alvo",        strategic.get("publico_alvo")        or "—"),
        ("Diferenciais",        strategic.get("diferenciais")        or "—"),
        ("Pilares Estratégicos",strategic.get("pilares")             or "—"),
    ]
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(fields):
        row             = tbl.rows[i]
        row.cells[0].text = lbl
        row.cells[1].text = val
        r0 = row.cells[0].paragraphs[0].runs[0]
        r0.font.bold = True; r0.font.color.rgb = C_BLUE
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)

    if strategic.get("objetivos_estrategicos"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Objetivos Estratégicos"); r.font.bold = True; r.font.color.rgb = C_BLUE
        doc.add_paragraph(strategic["objetivos_estrategicos"])

    doc.add_page_break()

    # ── Dashboard de KPIs — tabela resumo ───────────────────
    if okrs:
        h1 = doc.add_heading("Dashboard de KPIs — Resumo", 1)
        for run in h1.runs: run.font.color.rgb = C_BLUE

        headers = ["KPI", "Área", "Un.", "Total Planejado",
                   "Total Realizado", "% Realização", "Meses Preench."]
        tbl = doc.add_table(rows=1 + len(okrs), cols=len(headers))
        tbl.style = "Table Grid"
        _header_row(tbl, headers)

        for i, o in enumerate(okrs):
            o   = _ensure_okr_meses(o)
            tp  = sum(float(m.get("previsto",  0)) for m in o["meses"])
            tr  = sum(float(m.get("realizado", 0)) for m in o["meses"])
            pct = (tr / tp * 100) if tp > 0 else 0
            sem = "Verde" if pct >= 95 else ("Amarelo" if pct >= 70 else "Vermelho")
            filled = sum(1 for m in o["meses"] if float(m.get("realizado", 0)) != 0)
            row = tbl.rows[i + 1]
            row.cells[0].text = o.get("nome", "")
            row.cells[1].text = o.get("area", "")
            row.cells[2].text = o.get("unidade", "")
            row.cells[3].text = f"{tp:.2f}"
            row.cells[4].text = f"{tr:.2f}"
            row.cells[5].text = f"{pct:.1f}%"
            row.cells[6].text = f"{filled}/36"
            r5 = row.cells[5].paragraphs[0].runs[0]
            r5.font.bold = True; r5.font.color.rgb = _color_pct(pct)

        doc.add_page_break()

        # ── KPIs individuais com gráficos ─────────────────────────
        h1 = doc.add_heading("KPIs — Análise Detalhada", 1)
        for run in h1.runs: run.font.color.rgb = C_BLUE

        for o in okrs:
            o   = _ensure_okr_meses(o)
            tp  = sum(float(m.get("previsto",  0)) for m in o["meses"])
            tr  = sum(float(m.get("realizado", 0)) for m in o["meses"])
            pct = (tr / tp * 100) if tp > 0 else 0
            sem = "🟢" if pct >= 95 else ("🟡" if pct >= 70 else "🔴")

            h2 = doc.add_heading(f"{sem} {o.get('nome','')} ({o.get('unidade','')})", 2)
            for run in h2.runs: run.font.color.rgb = C_BLUE

            # Linha de estatísticas
            sp = doc.add_paragraph()
            r  = sp.add_run(f"Área: {o.get('area','—')}  |  "); r.font.color.rgb = C_GRAY
            r  = sp.add_run(f"Planejado: {tp:.2f}  |  "); r.font.color.rgb = C_BLUE
            r  = sp.add_run(f"Realizado: {tr:.2f}  |  "); r.font.color.rgb = _color_pct(pct)
            r  = sp.add_run(f"Realização: {pct:.1f}%"); r.font.bold = True; r.font.color.rgb = _color_pct(pct)

            # Gráfico
            try:
                img_bytes = _create_kpi_chart_image(o)
                doc.add_picture(io.BytesIO(img_bytes), width=Cm(16))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as exc:
                doc.add_paragraph(f"[Gráfico indisponível: {exc}]")

            doc.add_paragraph()  # espaçamento

        doc.add_page_break()

    # ── SWOT ─────────────────────────────────────────────────
    if swot_list:
        h1 = doc.add_heading("Análise SWOT", 1)
        for run in h1.runs: run.font.color.rgb = C_BLUE

        SWOT_FG = {"Força": RGBColor(6,95,70), "Fraqueza": RGBColor(153,27,27),
                   "Oportunidade": RGBColor(30,58,138), "Ameaça": RGBColor(146,64,14)}
        tbl = doc.add_table(rows=1 + len(swot_list), cols=3)
        tbl.style = "Table Grid"
        _header_row(tbl, ["Tipo", "Descrição", "Prioridade"])
        for i, item in enumerate(swot_list):
            tipo = item.get("tipo", "")
            row  = tbl.rows[i + 1]
            row.cells[0].text = tipo
            row.cells[1].text = item.get("descricao", "")
            row.cells[2].text = item.get("prioridade", "")
            r0 = row.cells[0].paragraphs[0].runs[0]
            r0.font.bold = True
            r0.font.color.rgb = SWOT_FG.get(tipo, C_GRAY)

        doc.add_page_break()

    # ── Planos de Ação ────────────────────────────────────────
    h1 = doc.add_heading("Planos de Ação", 1)
    for run in h1.runs: run.font.color.rgb = C_BLUE

    if actions:
        today  = date.today()
        STATUS_RBG = {
            "Concluído":    C_GREEN,
            "Em andamento": C_ORG,
            "Pendente":     C_GRAY,
            "Atrasado":     C_RED,
        }
        # TODAS as colunas para leitura completa pelo gestor
        headers = [
            "#", "Título", "Descrição", "KPI Vinculada", "Área",
            "Responsável", "D. Início", "Vencimento",
            "Como Fazer / Obs.", "Prioridade", "Status"
        ]
        tbl = doc.add_table(rows=1 + len(actions), cols=len(headers))
        tbl.style = "Table Grid"
        _header_row(tbl, headers)

        # Larguras sugeridas (soma ~17cm)
        col_widths = [Cm(0.7), Cm(3.0), Cm(2.8), Cm(2.2), Cm(1.8),
                      Cm(2.2), Cm(1.5), Cm(1.5), Cm(3.0), Cm(1.5), Cm(1.6)]
        for j, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[j].width = w

        for i, a in enumerate(actions):
            status    = a.get("status", "Pendente")
            dv_str    = a.get("data_vencimento", "") or ""
            di_str    = a.get("data_inicio",     "") or ""
            atrasado  = (
                status != "Concluído"
                and _safe_date(a.get("data_vencimento"))
                and _safe_date(a.get("data_vencimento")) < today
            )
            # Como Fazer + Observações juntos
            cf_obs = " | ".join(filter(None, [
                a.get("como_fazer",  ""),
                a.get("observacoes", ""),
            ])) or "—"

            row = tbl.rows[i + 1]
            row.cells[0].text  = str(i + 1)
            row.cells[1].text  = a.get("titulo",      "—") or "—"
            row.cells[2].text  = a.get("descricao",   "—") or "—"
            row.cells[3].text  = a.get("okr",         "—") or "—"
            row.cells[4].text  = a.get("area",        "—") or "—"
            row.cells[5].text  = a.get("responsavel", "—") or "—"
            row.cells[6].text  = di_str or "—"

            dv_cell = row.cells[7]
            dv_cell.text = f"{dv_str} (!)" if atrasado else (dv_str or "—")
            if atrasado:
                for r in dv_cell.paragraphs[0].runs:
                    r.font.bold = True; r.font.color.rgb = C_RED

            row.cells[8].text  = cf_obs
            row.cells[9].text  = a.get("prioridade", "—") or "—"
            sc = row.cells[10]
            sc.text = status
            for r in sc.paragraphs[0].runs:
                r.font.bold = True
                r.font.color.rgb = STATUS_RBG.get(status, C_GRAY)
    else:
        doc.add_paragraph("Nenhum plano de ação cadastrado.")

    # ── Rodapé da última página ──────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph(
        f"BK Engenharia e Tecnologia  |  Planejamento Estratégico  |  {today_str}"
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs:
        r.font.size = Pt(9); r.font.color.rgb = C_GRAY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_html_report(dados: dict) -> str:
    today_str  = date.today().strftime("%d/%m/%Y")
    strategic  = dados.get("strategic", {})
    okrs       = dados.get("s", [])
    actions    = dados.get("actions", [])
    swot       = dados.get("swot",    [])

    total_prev = sum(float(m.get("previsto",  0)) for o in okrs for m in _ensure_okr_meses(o)["meses"])
    total_real = sum(float(m.get("realizado", 0)) for o in okrs for m in _ensure_okr_meses(o)["meses"])
    pct_geral  = (total_real / total_prev * 100) if total_prev > 0 else 0
    n_concluidos = sum(1 for a in actions if a.get("status") == "Concluído")
    n_atrasados  = sum(
        1 for a in actions
        if a.get("status") != "Concluído"
        and _safe_date(a.get("data_vencimento"))
        and _safe_date(a.get("data_vencimento")) < date.today()
    )
    cor_geral = "#059669" if pct_geral >= 95 else ("#D97706" if pct_geral >= 70 else "#DC2626")

    overview_html = ""
    if okrs:
        ov_names, ov_prevs, ov_reals, ov_pcts = [], [], [], []
        for o in okrs:
            o2  = _ensure_okr_meses(o)
            tp  = sum(float(m.get("previsto",  0)) for m in o2["meses"])
            tr  = sum(float(m.get("realizado", 0)) for m in o2["meses"])
            pct = (tr / tp * 100) if tp > 0 else 0
            ov_names.append(o2.get("nome", "")[:28])
            ov_prevs.append(tp); ov_reals.append(tr); ov_pcts.append(round(pct, 1))
        fig_ov = go.Figure()
        fig_ov.add_trace(go.Bar(name="Planejado", x=ov_names, y=ov_prevs,
                                marker_color=BK_BLUE_LIGHT, opacity=0.75))
        fig_ov.add_trace(go.Bar(name="Realizado", x=ov_names, y=ov_reals,
                                marker_color=BK_GREEN))
        fig_ov.add_trace(go.Scatter(name="% Realização", x=ov_names, y=ov_pcts,
                                    mode="markers+text", yaxis="y2",
                                    marker=dict(size=10, color=BK_ORANGE),
                                    text=[f"{p:.0f}%" for p in ov_pcts],
                                    textposition="top center",
                                    textfont=dict(size=9)))
        fig_ov.update_layout(barmode="group", height=420,
                             yaxis2=dict(overlaying="y", side="right",
                                         title="% Realização", range=[0, 160]),
                             paper_bgcolor="white", plot_bgcolor="#F8FAFC",
                             margin=dict(l=50, r=60, t=50, b=130),
                             legend=dict(orientation="h", y=1.05),
                             font=dict(family="Segoe UI"),
                             xaxis=dict(tickangle=-45, tickfont=dict(size=9)))
        overview_html = fig_ov.to_html(full_html=False, include_plotlyjs=False)

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<title>Relatório Planejamento Estratégico — BK Engenharia</title>
<script src="https://cdn.plot.ly/plotly-2.27.0.min.js"></script>
<style>
  body{{font-family:'Segoe UI',sans-serif;background:#F0F4F8;margin:0;padding:20px;color:#1a202c}}
  .hero{{background:linear-gradient(135deg,#1565C0,#00897B);color:white;padding:32px;border-radius:12px;margin-bottom:24px}}
  .hero h1{{margin:0;font-size:24px}}.hero p{{margin:6px 0 0;opacity:.85;font-size:13px}}
  .card{{background:white;border-radius:10px;padding:20px 24px;margin-bottom:20px;box-shadow:0 2px 8px rgba(0,0,0,.06)}}
  .card h2{{font-size:14px;color:#1565C0;border-bottom:2px solid #E3F2FD;padding-bottom:8px;margin-top:0}}
  table{{width:100%;border-collapse:collapse;font-size:12px}}
  th{{background:#1565C0;color:white;padding:7px 10px;text-align:left}}
  td{{padding:6px 10px;border-bottom:1px solid #E2E8F0}}
  tr:nth-child(even) td{{background:#F8FAFC}}
</style>
</head>
<body>
<div class="hero"><h1>📊 Planejamento Estratégico — BK Engenharia e Tecnologia</h1>
<p>Gerado em {today_str} &nbsp;|&nbsp; Horizonte: 36 meses</p></div>
<div class="card"><h2>🧭 Norte Estratégico</h2>
<table><tbody>
<tr><td><b>Visão</b></td><td>{strategic.get('visao') or '—'}</td></tr>
<tr><td><b>Missão</b></td><td>{strategic.get('missao') or '—'}</td></tr>
<tr><td><b>Proposta de Valor</b></td><td>{strategic.get('proposta_valor') or '—'}</td></tr>
<tr><td><b>Público-Alvo</b></td><td>{strategic.get('publico_alvo') or '—'}</td></tr>
</tbody></table></div>
{f'<div class="card"><h2>📊 KPIs</h2>{overview_html}</div>' if overview_html else ''}
</body></html>"""
    return html


# ─── Views Django ─────────────────────────────────────────────────────────────

@login_required
def dashboard(request):
    dados   = get_planning()
    today   = date.today()
    okrs    = dados.get("s", [])
    actions = dados.get("actions", [])

    total_prev = sum(float(m.get("previsto",  0)) for o in okrs for m in _ensure_okr_meses(o)["meses"])
    total_real = sum(float(m.get("realizado", 0)) for o in okrs for m in _ensure_okr_meses(o)["meses"])
    pct_real   = (total_real / total_prev * 100) if total_prev > 0 else 0

    n_atrasados  = sum(1 for a in actions if a.get("status") != "Concluído"
                       and _safe_date(a.get("data_vencimento"))
                       and _safe_date(a.get("data_vencimento")) < today)
    n_concluidos = sum(1 for a in actions if a.get("status") == "Concluído")
    n_andamento  = sum(1 for a in actions if a.get("status") == "Em andamento")

    atrasados = []
    for a in actions:
        dv = _safe_date(a.get("data_vencimento"))
        if a.get("status") != "Concluído" and dv and dv < today:
            atrasados.append({**a, "dias_atraso": (today - dv).days})
    atrasados.sort(key=lambda x: x["dias_atraso"], reverse=True)

    return render(request, "planejamento/dashboard.html", {
        "dados":            dados,
        "n_okrs":           len(okrs),
        "pct_real":         round(pct_real, 1),
        "pct_real_color":   BK_GREEN if pct_real >= 90 else (BK_ORANGE if pct_real >= 70 else BK_RED),
        "n_actions":        len(actions),
        "n_concluidos":     n_concluidos,
        "n_atrasados":      n_atrasados,
        "n_andamento":      n_andamento,
        "fig_overview_json":fig_okrs_overview(dados),
        "fig_status_json":  fig_actions_status(dados),
        "fig_swot_json":    fig_swot_quadrant(dados.get("swot", [])) if dados.get("swot") else None,
        "gauges_json":      [fig_okr_gauge(o) for o in okrs[:4]],
        "atrasados":        atrasados[:10],
    })


@login_required
def socios(request):
    dados = get_planning()
    if request.method == "POST":
        action = request.POST.get("action")
        if action == "add":
            nome = request.POST.get("nome", "").strip()
            if nome:
                dados["partners"].append({
                    "nome": nome,
                    "cargo": request.POST.get("cargo", ""),
                    "email": request.POST.get("email", ""),
                    "telefone": request.POST.get("telefone", ""),
                    "observacoes": request.POST.get("observacoes", ""),
                })
                save_planning(dados)
                messages.success(request, "Sócio/Gestor adicionado!")
            else:
                messages.warning(request, "Informe o nome.")

        elif action == "save_table":
            rows_json = request.POST.get("rows_json", "[]")
            try:
                rows = json.loads(rows_json)
                dados["partners"] = _normalize_partner_rows(rows)
                save_planning(dados)
                messages.success(request, "Sócios/Gestores salvos!")
            except Exception as e:
                messages.error(request, f"Erro: {e}")

        return redirect("planejamento:socios")
    return render(request, "planejamento/socios.html", {"dados": dados})


@login_required
def estrategia(request):
    dados = get_planning()
    if request.method == "POST":
        for f in ["visao","missao","valores","posicionamento","proposta_valor",
                  "publico_alvo","diferenciais","pilares","objetivos_estrategicos","notas"]:
            dados["strategic"][f] = request.POST.get(f, "")
        save_planning(dados)
        messages.success(request, "Estratégia salva!")
        return redirect("planejamento:estrategia")
    return render(request, "planejamento/estrategia.html", {"dados": dados})


@login_required
def areas(request):
    dados = get_planning()
    if request.method == "POST":
        action = request.POST.get("action")
        if action == "add":
            area = request.POST.get("area", "").strip()
            if area:
                dados["areas"].append({
                    "area": area,
                    "responsavel": request.POST.get("responsavel", ""),
                    "email": request.POST.get("email", ""),
                    "observacoes": request.POST.get("observacoes", ""),
                })
                save_planning(dados)
                messages.success(request, "Área adicionada!")
            else:
                messages.warning(request, "Informe a área.")

        elif action == "save_table":
            rows_json = request.POST.get("rows_json", "[]")
            try:
                rows = json.loads(rows_json)
                dados["areas"] = _normalize_area_rows(rows)
                save_planning(dados)
                messages.success(request, "Áreas salvas!")
            except Exception as e:
                messages.error(request, f"Erro: {e}")

        return redirect("planejamento:areas")
    return render(request, "planejamento/areas.html", {"dados": dados})


@login_required
def swot(request):
    dados = get_planning()
    if request.method == "POST":
        if request.POST.get("action") == "save_table":
            rows_json = request.POST.get("rows_json", "[]")
            try:
                rows = json.loads(rows_json)
                dados["swot"] = [
                    {
                        "tipo":      _clean_text(r.get("tipo")),
                        "descricao": _clean_text(r.get("descricao")),
                        "prioridade":_clean_text(r.get("prioridade"), "Média") or "Média",
                    }
                    for r in rows
                    if _clean_text(r.get("descricao")) and not r.get("excluir", False)
                ]
                save_planning(dados)
                messages.success(request, "SWOT salva!")
            except Exception as e:
                messages.error(request, f"Erro: {e}")
        return redirect("planejamento:swot")

    fig_swot_json = fig_swot_quadrant(dados.get("swot", [])) if dados.get("swot") else None
    swot_groups   = _group_swot_items(dados.get("swot", []))
    return render(request, "planejamento/swot.html", {
        "dados": dados,
        "fig_swot_json": fig_swot_json,
        "swot_groups":   swot_groups,
        "tipos":         ["Força", "Fraqueza", "Oportunidade", "Ameaça"],
        "prioridades":   ["Alta", "Média", "Baixa"],
    })


@login_required
def kpis(request):
    """
    View de KPIs (renomeada de 's' para 'kpis').
    BUG FIX principal: save_meta agora faz lookup por NOME, não por índice.
    """
    dados = get_planning()
    dados.setdefault("s", [])
    dados["s"] = [_ensure_okr_meses(o) for o in dados["s"]]

    unidade_opts = ["R$", "%", "un", "clientes", "projetos", "h", "dias", "índice"]
    month_cols   = [f"M{i:02d}" for i in range(1, 37)]

    if request.method == "POST":
        action = request.POST.get("action", "").strip()

        # ── Salvar metadados dos KPIs ───────────────────────────────────────
        if action == "save_meta":
            rows_json = request.POST.get("rows_json", "[]")
            try:
                rows = json.loads(rows_json)

                # BUG FIX: indexar por NOME para preservar dados ao editar/reordenar
                antigos_by_nome = {}
                for o in dados.get("s", []):
                    nome_key = _clean_text(o.get("nome", ""))
                    if nome_key:
                        antigos_by_nome[nome_key] = _ensure_okr_meses(dict(o))

                novos = []
                for row in rows:
                    if row.get("excluir"):
                        continue
                    nome = _clean_text(row.get("nome"))
                    if not nome:
                        continue

                    # Busca dados existentes pelo NOME (não pelo índice)
                    existente = antigos_by_nome.get(nome)
                    meses     = existente.get("meses", []) if existente else []
                    while len(meses) < 36:
                        meses.append({"previsto": 0.0, "realizado": 0.0})

                    novos.append({
                        "nome":     nome,
                        "area":     _clean_text(row.get("area")),
                        "unidade":  _clean_text(row.get("unidade"), "un"),
                        "descricao":_clean_text(row.get("descricao")),
                        "inicio":   _clean_text(row.get("inicio")),
                        "meses":    meses[:36],
                    })

                dados["s"] = novos
                save_planning(dados)
                messages.success(request, "KPIs salvos com sucesso.")
            except Exception as e:
                messages.error(request, f"Erro ao salvar KPIs: {e}")
            return redirect("planejamento:okrs")

        # ── Salvar Planejado ou Realizado (AJAX) ────────────────────────────
        elif action in ("save_previsto", "save_realizado"):
            campo   = "previsto" if action == "save_previsto" else "realizado"
            is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"
            rows_json = request.POST.get("rows_json", "[]")
            try:
                rows = json.loads(rows_json)
                if not isinstance(rows, list):
                    raise ValueError("rows_json deve ser uma lista")

                s_list = dados.get("s", [])
                for row in rows:
                    if not isinstance(row, dict):
                        continue
                    try:
                        idx = int(row.get("idx", -1))
                    except (TypeError, ValueError):
                        continue
                    if idx < 0 or idx >= len(s_list):
                        continue
                    okr   = s_list[idx]
                    if not isinstance(okr, dict):
                        continue
                    meses = okr.setdefault("meses", [])
                    while len(meses) < 36:
                        meses.append({"previsto": 0.0, "realizado": 0.0})
                    for i in range(36):
                        key = f"M{i+1:02d}"
                        if key in row:
                            try:
                                meses[i][campo] = float(row[key] or 0)
                            except (TypeError, ValueError):
                                pass

                save_planning(dados)
                label = "Planejado" if campo == "previsto" else "Realizado"

                if is_ajax:
                    non_zero = sum(1 for okr in dados.get("s", [])
                                   for m in okr.get("meses", [])
                                   if (m.get(campo) or 0) != 0)
                    try:
                        db_obj   = PlanningData.objects.get(slug="bk")
                        verify_nz = sum(1 for okr in db_obj.dados.get("s", [])
                                        for m in okr.get("meses", [])
                                        if (m.get(campo) or 0) != 0)
                    except Exception:
                        verify_nz = -1
                    return JsonResponse({
                        "status":    "ok",
                        "msg":       f"{label} salvo com sucesso.",
                        "non_zero":  non_zero,
                        "verify_db": verify_nz,
                    })
                messages.success(request, f"{label} salvo com sucesso.")
            except Exception as e:
                import traceback; traceback.print_exc()
                if is_ajax:
                    return JsonResponse({"status": "error", "msg": str(e)}, status=400)
                messages.error(request, f"Erro ao salvar: {e}")

            return redirect("planejamento:okrs")

    s_list = [_ensure_okr_meses(dict(o)) for o in dados.get("s", [])]
    return render(request, "planejamento/okrs.html", {
        "dados":            dados,
        "s_list":           s_list,
        "okrs_list":        s_list,   # alias único (sem duplicação)
        "unidade_opts":     unidade_opts,
        "month_cols":       month_cols,
        "fig_overview_json":fig_okrs_overview(dados),
    })


@login_required
def kpi_detail_json(request, nome):
    dados = get_planning()
    okr   = None
    for item in dados.get("s", []):
        if str(item.get("nome", "")).strip() == str(nome).strip():
            okr = _ensure_okr_meses(item)
            break
    if not okr:
        return JsonResponse({"error": "KPI não encontrada."}, status=404)

    tp  = sum(float(m.get("previsto",  0) or 0) for m in okr["meses"])
    tr  = sum(float(m.get("realizado", 0) or 0) for m in okr["meses"])
    pct = round((tr / tp * 100), 1) if tp > 0 else 0.0

    labels = _month_labels_for_okr(okr)
    table  = []
    for i, mes in enumerate(okr["meses"]):
        prev = float(mes.get("previsto",  0) or 0)
        real = float(mes.get("realizado", 0) or 0)
        diff = real - prev
        status_val = ("Acima/atingido" if real >= prev and prev > 0
                      else ("Abaixo" if real > 0 and real < prev else "Sem realização"))
        table.append({"mes": labels[i], "prev": prev, "real": real,
                      "diff": diff, "status": status_val})

    return JsonResponse({
        "nome":           okr.get("nome", ""),
        "unidade":        okr.get("unidade", ""),
        "tp":             tp, "tr": tr, "pct": pct,
        "fig_gauge":      fig_okr_gauge(okr),
        "fig_monthly":    fig_okr_monthly(okr),
        "fig_cumulative": fig_okr_cumulative(okr),
        "table":          table,
    })


@login_required
def planos_acao(request):
    dados = get_planning()
    today = date.today()

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "add":
            titulo = request.POST.get("titulo", "").strip()
            if not titulo:
                messages.warning(request, "Informe o título.")
            else:
                dados["actions"].append({
                    "titulo":          titulo,
                    "area":            request.POST.get("area",            ""),
                    "responsavel":     request.POST.get("responsavel",     ""),
                    "okr":             request.POST.get("okr",             ""),
                    "descricao":       request.POST.get("descricao",       ""),
                    "data_inicio":     request.POST.get("data_inicio",     ""),
                    "data_vencimento": request.POST.get("data_vencimento", ""),
                    "status":          request.POST.get("status",          "Pendente"),
                    "observacoes":     request.POST.get("observacoes",     ""),
                    "como_fazer":      request.POST.get("como_fazer",      ""),
                    "prioridade":      request.POST.get("prioridade",      "Média"),
                })
                save_planning(dados)
                messages.success(request, f'Plano "{titulo}" adicionado!')

        elif action == "save_table":
            rows_json = request.POST.get("rows_json", "[]")
            try:
                rows = json.loads(rows_json)
                dados["actions"] = [
                    r for r in rows
                    if r.get("titulo", "").strip() and not r.get("excluir", False)
                ]
                save_planning(dados)
                messages.success(request, "Planos salvos!")
            except Exception as e:
                messages.error(request, f"Erro: {e}")

        return redirect("planejamento:planos_acao")

    n_total      = len(dados["actions"])
    n_concluidos = sum(1 for a in dados["actions"] if a.get("status") == "Concluído")
    n_andamento  = sum(1 for a in dados["actions"] if a.get("status") == "Em andamento")
    n_atrasados  = sum(1 for a in dados["actions"]
                       if a.get("status") != "Concluído"
                       and _safe_date(a.get("data_vencimento"))
                       and _safe_date(a.get("data_vencimento")) < today)

    fig_resp_json = None
    if dados["actions"]:
        resp_atraso = {}
        for a in dados["actions"]:
            dv = _safe_date(a.get("data_vencimento"))
            if dv and a.get("status") != "Concluído" and dv < today:
                r = a.get("responsavel", "N/A")
                resp_atraso[r] = resp_atraso.get(r, 0) + (today - dv).days
        if resp_atraso:
            df_r = pd.DataFrame(list(resp_atraso.items()), columns=["Responsável", "Atraso"])
            fig_r = px.bar(df_r, x="Responsável", y="Atraso",
                           title="Atraso total por Responsável (dias)",
                           color="Atraso", color_continuous_scale=[BK_ORANGE, BK_RED])
            _fig_layout(fig_r, height=320)
            fig_resp_json = fig_r.to_json()

    return render(request, "planejamento/planos_acao.html", {
        "dados":           dados,
        "today":           today.isoformat(),
        "n_total":         n_total,
        "n_concluidos":    n_concluidos,
        "n_andamento":     n_andamento,
        "n_atrasados":     n_atrasados,
        "fig_status_json": fig_actions_status(dados),
        "fig_timeline_json": fig_actions_timeline(dados),
        "fig_resp_json":   fig_resp_json,
        "status_opts":     ["Pendente", "Em andamento", "Concluído"],
    })


@login_required
def relatorios(request):
    dados   = get_planning()
    today   = date.today()
    okrs    = dados.get("s",       [])
    actions = dados.get("actions", [])
    swot    = dados.get("swot",    [])

    n_atrasados  = sum(1 for a in actions if a.get("status") != "Concluído"
                       and _safe_date(a.get("data_vencimento"))
                       and _safe_date(a.get("data_vencimento")) < today)
    n_concluidos = sum(1 for a in actions if a.get("status") == "Concluído")
    n_andamento  = sum(1 for a in actions if a.get("status") == "Em andamento")
    n_pendente   = sum(1 for a in actions if a.get("status") == "Pendente")
    n_total_ac   = len(actions)

    total_prev_geral = 0.0
    total_real_geral = 0.0
    saude      = []
    kpi_charts = []

    for o in okrs:
        o      = _ensure_okr_meses(o)
        tp     = sum(float(m.get("previsto",  0)) for m in o["meses"])
        tr     = sum(float(m.get("realizado", 0)) for m in o["meses"])
        total_prev_geral += tp
        total_real_geral += tr
        pct    = (tr / tp * 100) if tp > 0 else 0
        filled = sum(1 for m in o["meses"] if float(m.get("realizado", 0)) != 0)
        semaforo = "🟢" if pct >= 95 else ("🟡" if pct >= 70 else "🔴")
        color    = BK_GREEN if pct >= 95 else (BK_ORANGE if pct >= 70 else BK_RED)

        saude.append({"semaforo": semaforo, "nome": o.get("nome"), "area": o.get("area"),
                      "unidade": o.get("unidade"), "pct": round(pct, 1),
                      "tp": round(tp, 2), "tr": round(tr, 2),
                      "filled": filled, "color": color})

        labels = _month_labels_for_okr(o)
        prevs  = [float(m.get("previsto",  0)) for m in o["meses"]]
        reals  = [float(m.get("realizado", 0)) for m in o["meses"]]

        fig_bar = go.Figure(data=[
            go.Bar(name="Planejado", x=labels, y=prevs,
                   marker_color=BK_BLUE_LIGHT, opacity=0.85),
            go.Bar(name="Realizado", x=labels, y=reals, marker_color=BK_GREEN),
        ])
        fig_bar.update_layout(barmode="group", height=240,
                              margin=dict(l=40, r=10, t=16, b=50),
                              legend=dict(orientation="h", y=1.08),
                              paper_bgcolor="white", plot_bgcolor="white",
                              xaxis=dict(tickangle=-45, tickfont=dict(size=9)))

        gauge_color = BK_GREEN if pct >= 95 else (BK_ORANGE if pct >= 70 else BK_RED)
        fig_gauge = go.Figure(go.Indicator(
            mode="gauge+number", value=round(pct, 1),
            number={"suffix": "%", "font": {"size": 22, "color": gauge_color}},
            gauge={"axis": {"range": [0, 120]}, "bar": {"color": gauge_color},
                   "steps": [{"range":[0,70],"color":"#FEE2E2"},
                              {"range":[70,95],"color":"#FEF3C7"},
                              {"range":[95,120],"color":"#D1FAE5"}],
                   "threshold": {"line":{"color":"#1565C0","width":2},
                                 "thickness":0.75,"value":100}},
        ))
        fig_gauge.update_layout(height=180, margin=dict(l=10,r=10,t=10,b=10),
                                paper_bgcolor="white")

        kpi_charts.append({
            "nome": o.get("nome"), "area": o.get("area"),
            "unidade": o.get("unidade"), "tp": round(tp,2), "tr": round(tr,2),
            "pct": round(pct,1), "semaforo": semaforo, "color": color,
            "fig_bar": fig_bar.to_json(), "fig_gauge": fig_gauge.to_json(),
        })

    pct_geral = (total_real_geral / total_prev_geral * 100) if total_prev_geral > 0 else 0

    by_status = {}
    by_area   = {}
    for a in actions:
        status_val = a.get("status", "Pendente")
        area_val   = a.get("area",   "—")
        by_status[status_val] = by_status.get(status_val, 0) + 1
        by_area[area_val]     = by_area.get(area_val,     0) + 1

    fig_ac_status_json = ""
    if by_status:
        fig_ac = go.Figure(go.Pie(
            labels=list(by_status.keys()), values=list(by_status.values()),
            marker_colors=[STATUS_COLORS.get(k, BK_GRAY) for k in by_status],
            hole=0.4, textinfo="label+percent", textfont_size=11,
        ))
        fig_ac.update_layout(height=260, margin=dict(l=10,r=10,t=10,b=10),
                             paper_bgcolor="white", showlegend=False)
        fig_ac_status_json = fig_ac.to_json()

    fig_ac_area_json = ""
    if by_area:
        sorted_areas = sorted(by_area.items(), key=lambda x: x[1], reverse=True)
        fig_area = go.Figure(go.Bar(
            x=[i[0] for i in sorted_areas], y=[i[1] for i in sorted_areas],
            marker_color=BK_BLUE,
            text=[i[1] for i in sorted_areas], textposition="outside",
        ))
        fig_area.update_layout(height=260, margin=dict(l=10,r=10,t=16,b=60),
                               paper_bgcolor="white", plot_bgcolor="white",
                               xaxis=dict(tickangle=-30, tickfont=dict(size=10)),
                               yaxis=dict(showgrid=True, gridcolor="#E2E8F0"))
        fig_ac_area_json = fig_area.to_json()

    recs = []
    threats    = [s for s in swot if s.get("tipo") == "Ameaça"       and s.get("prioridade") == "Alta"]
    opps       = [s for s in swot if s.get("tipo") == "Oportunidade" and s.get("prioridade") == "Alta"]
    weaknesses = [s for s in swot if s.get("tipo") == "Fraqueza"     and s.get("prioridade") == "Alta"]
    if threats:    recs.append(f"🔴 {len(threats)} Ameaça(s) Alta — crie planos de mitigação imediatos.")
    if opps:       recs.append(f"🔵 {len(opps)} Oportunidade(s) Alta — transforme em KPIs estratégicos.")
    if weaknesses: recs.append(f"🟡 {len(weaknesses)} Fraqueza(s) Alta — endereçar com planos de ação de curto prazo.")
    if n_atrasados: recs.append(f"⚠️ {n_atrasados} plano(s) atrasado(s) — replaneje: escopo, capacidade, nova data.")
    if okrs:       recs.append("📅 Estabeleça revisão mensal do realizado e trimestral das prioridades.")
    low_fill = [o["nome"] for o in saude if o["filled"] < 3]
    if low_fill:   recs.append(f"📊 KPIs com pouco histórico: {', '.join(low_fill[:3])} — preencha o realizado mensalmente.")
    if not recs:   recs.append("✅ Preencha Visão/Missão, SWOT e KPIs para gerar recomendações automáticas.")

    return render(request, "planejamento/relatorios.html", {
        "dados":               dados,
        "saude":               saude,
        "recs":                recs,
        "total_prev_geral":    round(total_prev_geral, 2),
        "total_real_geral":    round(total_real_geral, 2),
        "pct_geral":           round(pct_geral, 1),
        "kpi_charts":          kpi_charts,
        "actions":             actions,
        "n_total_ac":          n_total_ac,
        "n_concluidos":        n_concluidos,
        "n_andamento":         n_andamento,
        "n_pendente":          n_pendente,
        "n_atrasados":         n_atrasados,
        "fig_ac_status_json":  fig_ac_status_json,
        "fig_ac_area_json":    fig_ac_area_json,
        "today_str":           today.isoformat(),
    })


# ─── Endpoints de exportação ─────────────────────────────────────────────────

@login_required
def export_json(request):
    dados   = get_planning()
    content = json.dumps(dados, ensure_ascii=False, indent=2)
    resp    = HttpResponse(content, content_type="application/json")
    resp["Content-Disposition"] = 'attachment; filename="planejamento_export.json"'
    return resp


@login_required
def export_excel_view(request):
    dados = get_planning()
    xlsx  = export_excel(dados)
    resp  = HttpResponse(xlsx, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    resp["Content-Disposition"] = 'attachment; filename="planejamento_completo.xlsx"'
    return resp


@login_required
def export_zip_view(request):
    dados = get_planning()
    z     = export_csv_zip(dados)
    resp  = HttpResponse(z, content_type="application/zip")
    resp["Content-Disposition"] = 'attachment; filename="planning_csvs.zip"'
    return resp


@login_required
def export_html_view(request):
    dados = get_planning()
    html  = build_html_report(dados)
    resp  = HttpResponse(html, content_type="text/html; charset=utf-8")
    resp["Content-Disposition"] = 'attachment; filename="relatorio_planejamento.html"'
    return resp


@login_required
def export_word_view(request):
    """Download do relatório Word (.docx) com gráficos matplotlib."""
    dados = get_planning()
    try:
        docx_bytes = build_word_report(dados)
        resp = HttpResponse(
            docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp["Content-Disposition"] = 'attachment; filename="relatorio_planejamento.docx"'
        return resp
    except ImportError:
        messages.error(request, "python-docx ou matplotlib não instalados. Execute: pip install python-docx matplotlib")
        return redirect("planejamento:relatorios")
    except Exception as e:
        import traceback; traceback.print_exc()
        messages.error(request, f"Erro ao gerar Word: {e}")
        return redirect("planejamento:relatorios")


@login_required
def import_json(request):
    if request.method == "POST" and request.FILES.get("json_file"):
        try:
            content = request.FILES["json_file"].read().decode("utf-8")
            dados   = json.loads(content)
            save_planning(dados)
            messages.success(request, "JSON importado e salvo!")
        except Exception as e:
            messages.error(request, f"Erro ao importar: {e}")
    return redirect("planejamento:dashboard")


# ─── Aliases de compatibilidade ──────────────────────────────────────────────
okrs            = kpis
okr_detail_json = kpi_detail_json
